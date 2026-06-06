#!/usr/bin/env python3
"""
Quest-On 사용의사확인서 (Letter of Intent to Use) - 1페이지 버전 v3
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── 상수 ──
FONT = "맑은 고딕"
PAGE_W = 21.0
MARGIN_LR = 2.3
USABLE_W = PAGE_W - MARGIN_LR * 2  # 16.4cm
LABEL_W = 2.8
VALUE_W = (USABLE_W - LABEL_W * 2) / 2  # 5.4cm
HEADER_BG = "EDF0F5"


def _font(run, size=9, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if color:
        run.font.color.rgb = color
    return run


def _shading(cell, color=HEADER_BG):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    )


def _valign(cell, val="center"):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{val}"/>')
    )


def _no_border(cell):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
    )


def _cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(size * 1.5)
    _font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def _para(doc, text, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          color=None, italic=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    _font(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def _set_table_widths(table, widths):
    """테이블 전체 행에 걸쳐 열 너비를 일관되게 적용"""
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w


def _label_cell(cell, text, size=9):
    """라벨 셀: 배경색 + 중앙정렬 + bold"""
    _shading(cell)
    _valign(cell)
    _cell_text(cell, text, bold=True, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)


def _value_cell(cell, text="", size=9):
    _valign(cell)
    _cell_text(cell, text, size=size)


# ═══════════════════════════════════════════════════════════
# 메인
# ═══════════════════════════════════════════════════════════
def create_loi_document():
    doc = Document()

    # ── 페이지 설정 ──
    sec = doc.sections[0]
    sec.page_width = Cm(PAGE_W)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(MARGIN_LR)
    sec.right_margin = Cm(MARGIN_LR)

    # ── 기본 스타일 ──
    s = doc.styles["Normal"]
    s.font.name = FONT
    s.font.size = Pt(9)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    s.paragraph_format.space_before = Pt(0)
    s.paragraph_format.space_after = Pt(0)

    # ═══════════════════════════════════════
    # 제목
    # ═══════════════════════════════════════
    _para(doc, "사 용 의 사 확 인 서",
          size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=1)
    _para(doc, "Letter of Intent to Use",
          size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x88, 0x88, 0x88), after=8)

    # ═══════════════════════════════════════
    # 1. 확인자 정보
    # ═══════════════════════════════════════
    _para(doc, "1. 확인자 정보", size=10.5, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=3)

    t1 = doc.add_table(rows=4, cols=4)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.style = "Table Grid"
    _set_table_widths(t1, [Cm(LABEL_W), Cm(VALUE_W), Cm(LABEL_W), Cm(VALUE_W)])

    rows_data_1 = [
        ("소속기관", "", "부서/학과", ""),
        ("성명", "", "직위/직급", ""),
        ("연락처", "", "이메일", ""),
        ("기관유형", "☐ 대학교   ☐ 연구기관   ☐ 공공기관   ☐ 기업   ☐ 기타(          )", "", ""),
    ]

    for i, (l1, v1, l2, v2) in enumerate(rows_data_1):
        r = t1.rows[i]
        _label_cell(r.cells[0], l1)
        if i == 3:
            r.cells[1].merge(r.cells[2]).merge(r.cells[3])
            _value_cell(r.cells[1], v1)
        else:
            _value_cell(r.cells[1], v1)
            _label_cell(r.cells[2], l2)
            _value_cell(r.cells[3], v2)

    # ═══════════════════════════════════════
    # 2. 사용 예정 서비스 정보
    # ═══════════════════════════════════════
    _para(doc, "2. 사용 예정 서비스 정보", size=10.5, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=3)

    t2 = doc.add_table(rows=4, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.style = "Table Grid"
    _set_table_widths(t2, [Cm(LABEL_W), Cm(VALUE_W), Cm(LABEL_W), Cm(VALUE_W)])

    rows_data_2 = [
        ("서비스명", "Quest-On (퀘스트온)", "서비스유형", "AI 기반 서술형 시험·평가 플랫폼"),
        ("사용 예정\n분야/과목", "", "예상 사용\n인원(학생 수)", ""),
        ("사용 예정\n시기", "20   년    월 ~ 20   년    월", "사용형태", "☐ 정규수업  ☐ 비교과\n☐ 시범운영  ☐ 기타"),
        ("사용목적", "", "", ""),
    ]

    for i, (l1, v1, l2, v2) in enumerate(rows_data_2):
        r = t2.rows[i]
        _label_cell(r.cells[0], l1)
        if i == 3:
            r.cells[1].merge(r.cells[2]).merge(r.cells[3])
            _value_cell(r.cells[1], v1)
            r.height = Cm(1.8)
        else:
            _value_cell(r.cells[1], v1)
            _label_cell(r.cells[2], l2)
            _value_cell(r.cells[3], v2)

    # ═══════════════════════════════════════
    # 3. 서비스 주요 기능 (참고)
    # ═══════════════════════════════════════
    _para(doc, "3. 서비스 주요 기능 (참고)", size=10.5, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=3)

    t3 = doc.add_table(rows=2, cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.style = "Table Grid"
    col_w3 = Cm(USABLE_W / 4)
    _set_table_widths(t3, [col_w3] * 4)

    features = [
        ("AI 대화형 시험",   "시험 중 AI에게 질문하며\n사고를 심화하는 평가 방식"),
        ("AI 자동 채점",     "루브릭 기반 AI가 채팅 과정과\n답안을 종합 평가"),
        ("RAG 기반 맥락 답변", "수업 자료 기반으로 AI가\n맥락에 맞는 답변 제공"),
        ("실시간 모니터링",   "학생별 응시 현황·AI 대화를\n교수자가 실시간 확인"),
    ]

    for j, (title, desc) in enumerate(features):
        _label_cell(t3.rows[0].cells[j], title, size=8.5)
        _valign(t3.rows[1].cells[j])
        _cell_text(t3.rows[1].cells[j], desc, size=8.5,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════════════════════════════════
    # 4. 사용 의사 확인
    # ═══════════════════════════════════════
    _para(doc, "4. 사용 의사 확인", size=10.5, bold=True,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=10, after=3)

    t4 = doc.add_table(rows=1, cols=1)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    t4.style = "Table Grid"
    t4.rows[0].cells[0].width = Cm(USABLE_W)
    cell = t4.rows[0].cells[0]

    body = (
        "본인은 위에 기재된 「Quest-On」 서비스의 주요 기능 및 활용 방안에 대하여 "
        "충분히 설명을 듣고 이해하였으며, 해당 서비스가 정식 출시될 경우 "
        "본인의 수업 및 평가 활동에 적극적으로 도입·활용할 의사가 있음을 확인합니다."
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(16)
    p.paragraph_format.first_line_indent = Cm(0.4)
    _font(p.add_run(body), size=9)

    note = (
        "※ 본 확인서는 서비스 사용에 대한 의향을 표명하는 것으로, "
        "법적 구속력이 있는 계약서가 아닙니다."
    )
    pn = cell.add_paragraph()
    pn.paragraph_format.space_before = Pt(3)
    pn.paragraph_format.space_after = Pt(4)
    pn.paragraph_format.line_spacing = Pt(13)
    _font(pn.add_run(note), size=8, color=RGBColor(0x77, 0x77, 0x77))

    # ═══════════════════════════════════════
    # 날짜 + 서명
    # ═══════════════════════════════════════
    _para(doc, "20      년        월        일",
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER, before=16, after=10)

    ts = doc.add_table(rows=3, cols=2)
    ts.alignment = WD_TABLE_ALIGNMENT.RIGHT
    for i, (lbl, val) in enumerate([
        ("소 속 :", ""),
        ("성 명 :", "                                   (서명 / 인)"),
        ("연락처 :", ""),
    ]):
        ts.rows[i].cells[0].width = Cm(1.8)
        ts.rows[i].cells[1].width = Cm(7.0)

        p_l = ts.rows[i].cells[0].paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_l.paragraph_format.space_before = Pt(1)
        p_l.paragraph_format.space_after = Pt(1)
        _font(p_l.add_run(lbl), size=10, bold=True)

        p_v = ts.rows[i].cells[1].paragraphs[0]
        p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_v.paragraph_format.space_before = Pt(1)
        p_v.paragraph_format.space_after = Pt(1)
        _font(p_v.add_run(val), size=10)

        for c in ts.rows[i].cells:
            _no_border(c)

    # ── 하단 수신처 ──
    _para(doc, "━" * 55, size=5, align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0xBB, 0xBB, 0xBB), before=10, after=2)
    _para(doc, "Quest-On  귀중", size=11, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=1, after=1)
    _para(doc, "문의: questonkr@gmail.com", size=8,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x99, 0x99, 0x99), before=0, after=0)

    # ── 저장 ──
    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       "Quest-On_사용의사확인서.docx")
    doc.save(out)
    print(f"✅ 생성 완료: {out}")


if __name__ == "__main__":
    create_loi_document()
